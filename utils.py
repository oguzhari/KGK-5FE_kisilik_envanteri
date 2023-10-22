import streamlit as st
from gsheetsdb import connect
import os
import smtplib
import requests as rs
import mimetypes
from email import encoders
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import random as rd

conn = connect()


def run_query(query):
    rows = conn.execute(query, headers=1)
    rows = rows.fetchall()
    return rows


def get_sheet():
    sheet_csv = st.secrets["public_sheet_csv"]
    res = rs.get(url=sheet_csv)
    open('google.csv', 'wb').write(res.content)
    content = pd.read_csv('google.csv', header=None)
    #rows = run_query(f'SELECT * FROM "{sheet_url}"')
    return content


def tanımla_analiz_et(ogrenci):
    envanter_analiz = ""
    arada_kalanlar = []
    envanter_degerlendirme_1_0 = ["Sosyal ortamlarda kaygı yaşayabilir.",
                                  "Sosyal ortamlardan ziyade daha sakin ortamları tercih etmeye yatkın olabilir.",
                                  "Yeni girdikleri sosyal çevrelerde sakin, sessiz, resmi ve ciddi bir duruş sergileyebilir.",
                                  "Olumsuz duygularını dışavurmaktansa içlerinde yaşamaya eğilimlidir olabilir."]
    envanter_degerlendirme_1_1 = ["Sosyal ortamlarda sıcak,ortama kolayca dahil olabilir.",
                                  "Konuşkan, neşeli, dinlerken ilgili görünen kişilik yapısına sahip olabilir.",
                                  "Enerji seviyesi yüksek olduğu için emir almaktan hoşlanmayabilir.",
                                  "Olumsuz tepkilerini içinde bastırmayabilir."]

    envanter_degerlendirme_2_0 = ["İnsanlara kolayca güvenebilen, takım çalışmasına kolayca uyum sağlayabilir.",
                                  "İnsanlarla çatışmadan ziyade uyumlu davranışlar davranışlar gösterme eğiliminde olabilir.",
                                  "İnsanlara güven seviyesi yüksektir ve takım çalışmasına uyumlu olabilir."]
    envanter_degerlendirme_2_1 = ["Genellikle çatışma ortamından beslenmeye yatkın olabilir.",
                                  "Her şeye şüphe ile yaklaşabilir.", "Olumsuz düşünmeye eğilimi yüksek olabilir.",
                                  "İkna etmeye çalışırken zorluk çıkartabilir.",
                                  "Çatışmalardan ve sonuçlarından çekinmeyebilir.",
                                  "Kendine yapılan kötü olayları hatırlamaya yatkın olabilir."]

    envanter_degerlendirme_3_0 = ["Gerçekleştirmesi gereken işleri ertelemeye eğilimli olabilir.",
                                  "Plan ve programa uyma eğilimi düşüktür olabilir.",
                                  "Motivasyonu iş sürecinde düşebilir."]
    envanter_degerlendirme_3_1 = ["Bir işe yoğunlaştığında detaylara dikkat etmeyebilir.",
                                  "Programa ve hedeflerine yönelik titiz olabilir.",
                                  "Düzenli, planlı davranma eğilimi yüksektir."]

    envanter_degerlendirme_4_0 = ["Olaylara gerçekçi bakış açısıyla yaklaşabilir.",
                                  "Duygu durumu sürekli değişmeyen huzurlu olabilir.",
                                  "Kaygı seviyesini duruma göre ayarlamaya yatkın olabilir."]
    envanter_degerlendirme_4_1 = ["Olumsuz düşünmeye eğilimli, kendini suçlamaya yatkın olabilir.",
                                  "Bilmediği durumlarda endişeye kapılma, kötümser düşünme ve gerilme eğilimi yüksek olabilir.",
                                  " Sorumluluk alırken başkalarının onayını fazla önemseyebilir."]

    envanter_degerlendirme_5_0 = ["Rutinleri seven, yeniliği tehlike gibi algılamaya eğilimli olabilir.",
                                  "Geleneksel ve alışılmış olana bağlılık seviyesi yüksektir olabilir."]
    envanter_degerlendirme_5_1 = ["Yenilik gereken alanları, fırsatlarını görme eğilimleri yüksek olabilir.",
                                  "Farklı değişik deneyimlere açık olma eğilimleri yüksektir.",
                                  " Farklı ilgi alanlarına eğilimli olabilir."]

    envanter_degerlendirme_6_0 = ["Yeni girdiği ortamda kendini ifade edebilir.",
                                  "İletişim becerisi yüksek, isteklerini dile getirirken zorlanmaya bilir.",
                                  "Olumlu düşünmeye eğilimleri yüksek olabilir."]
    envanter_degerlendirme_6_1 = ["Yeni girdiği ortamlarda konuşmak yerine dinlemeyi tercih edebilir.",
                                  "Duygularını dışavurmaktansa içinde yaşayabilir.",
                                  "Memnuniyetsizliğini dile getirme konusunda sıkıntı yaşayabilir."]

    envanter_degerlendirme_7_0 = ["Diğer insanlardan önce kendinilerini önemseyebilir.",
                                  "Karar alırken kendi avantajlarını ön planda tutabilir."]
    envanter_degerlendirme_7_1 = ["Merhametli, insanlara yardım etmekten keyif alabilir.",
                                  "İşbirliğine eğilimli ve insan ilişkileri güçlü olabilir."]

    envanter_degerlendirme_8_0 = ["Öz disiplin seviyesi yüksek olabilir.",
                                  "İçselkontrolü sağlayabilen, sorumluluklarını yerine getirmeye eğilimli olabilir.",
                                  "Başka insanlara karşı kendini sorumlu hissedebilir."]
    envanter_degerlendirme_8_1 = ["Öz disiplin seviyesi biraz düşük olabilir.",
                                  "Hedefine ulaşmakta zorluk yaşayabilir."]

    envanter_degerlendirme_9_0 = [
        "Karar vermede güçlük çekebilir. Sorumluluk almayı ve üstesinden gelmeyi sevmeyebilir. Konsantrasyon seviyeleri düşük olabilir.",
        "Endişe seviyeleri yüksektir."]
    envanter_degerlendirme_9_1 = ["Aldığı kararlardan emin, kendisi ve çevresi ile uyumlu olabilir.",
                                  "Problemler karşısında stres seviyelerini kontrol edebilmeye eğilimlidir."]

    envanter_degerlendirme_10_0 = ["Geleneksel görüşleri ve rutin işleri hayatlarında tercih edebilir.",
                                   "Riski sevmeyen, bildiği yolu bırakmayı tercih etmeyen kişiler olabilir."]
    envanter_degerlendirme_10_1 = ["Duygusal yaşamı renkli, farklı konularda düşünme eğilimi yüksek olabilir.",
                                   "İlgi alanı geniş olduğu için yeni bilgiler ile beslenebilir.",
                                   "Yeni işlere girişmekten hoşnut olabilir."]

    envanter_degerlendirme_11_0 = ["Sakindir, göze batmayan şekilde yaşamayı tercih edebilir.",
                                   "Çatışmaya girmekten kaçınabilir.",
                                   "Huzur hayatında vazgeçilmez bir kavram olabilir.",
                                   "Dürtüsel hareketlerde bulunmaktan kaçınabilir."]
    envanter_degerlendirme_11_1 = ["Hareket etmeyi seven, olumlu düşünmeye eğilimli olabilir.",
                                   "Konuşmayı ve aktiviteyi sevebilir.",
                                   "Davranışları sınırlandırıldığı ya da kendi istekleri doğrultusunda hareket edemedikleri zaman mutsuz olabilir."]

    envanter_degerlendirme_12_0 = ["Eleştiriye tahammül seviyesi yüksek olabilir.",
                                   "İnsanlara güvenmeye yatkın olabilir. Mantıksal çerçevede problemleri çözebilir.",
                                   "Birlikte olmaya, ekip çalışması yapmaya uyumludur."]
    envanter_degerlendirme_12_1 = ["İnsanlara güveme ihtimalleri düşük olabilir. Kolayca öfkelenebilir.",
                                   "Problemleri çözerken öfkeli olabilirler.",
                                   "Zaman zaman kaba olabilen, insanlarla empati kurmakta zorlanabilir.",
                                   "Çatışma ve sonuçlarından etkilenmeye bilir.",
                                   "Hakkı ararken çatışmaya girmekten çekinmeye bilir."]

    envanter_degerlendirme_13_0 = ["Motivasyon ve bir iş sürecini organizme etme olasılığı düşük olabilir.",
                                   "Yaptığı işten kolayca sıkılabilir.", "Zor olan işleri ertelemeye yatkın olabilir."]
    envanter_degerlendirme_13_1 = ["Amaca yönelik hareket edebilir.", "İçsel motivasyonu yüksek olabilir.",
                                   "Titizlikle sonuca odaklanma eğilimi yüksektir."]

    envanter_degerlendirme_14_0 = ["Davranışı önceden tahmin edilebilir.",
                                   "Duygu durumunu kontrol edebilen kişiler olabilir.",
                                   "İnsanlarla etkileşimi pozitif olmaya eğilimlidir."]
    envanter_degerlendirme_14_1 = ["Kolaylıkla üzülebilir, duygu durumu sık sık değişebilir.",
                                   "Bu yüzden de davranışları tutarsız olabilir.",
                                   "Gerilmekten dolayı motivasyonu düşük olabilir."]

    envanter_degerlendirme_15_0 = ["Yeni fikirlere açık değildir.", "Somut düşünmeye yatkın olabilir.",
                                   "Sabit yapıya sahip olabilir."]
    envanter_degerlendirme_15_1 = ["Olayları irdeleyebilir.", "Neden sonuç bağlantısı kurabilmeye yatkın olabilir.",
                                   "Bir olayın gelecek zamandaki avantaj ve dezavantajlarını önceden tahmin edebilir."]

    envanter_degerlendirme_16_0 = ["Sosyal ortamlarda girişkenlik göstermeyebilir.",
                                   "Fikirlerini söylemekten kaçınabilir.",
                                   "Başkalarıyla fikir alışverişine girmektense iç sesiyle konuşmaya eğilimlidir."]
    envanter_degerlendirme_16_1 = ["İletişim becerileri ve içsel motivasyonu yüksek olabilir.",
                                   "Sosyal becerileri gelişmiş, dominant kişilik yapısı olabilirler.",
                                   "Kendinden emin olmaya eğilimlidir."]

    envanter_degerlendirme_17_0 = ["Olumsuz düşüncelere ve şüpheci davranışlara eğilimli olabilir.",
                                   "Çalışma ortamlarında rekabetçi davranışlar sergileyebilir.",
                                   "Olumsuz durumlar karşısında fevri davranışlar sergileyebilir."]
    envanter_degerlendirme_17_1 = ["Sosyal ortamlarda çatışmalardan uzak kalmaya çalışır.",
                                   "Çalışma ortamlarında uyumlu olmaya eğilimlidir.",
                                   "Pozitif düşünmeye daha yatkındır."]

    envanter_degerlendirme_18_0 = [
        "Ayrıntılara önem veren  planlı ve programlı davranmaya yatkın olma eğilimleri yüksektir.",
        "Mükemmeliyetçi davranışlar sergileyebilir."]
    envanter_degerlendirme_18_1 = ["Yapılması gereken işleri ertelemeye yatkın olabilir.",
                                   "Plan ve programlarını uygulamaya eğilimleri düşük olabilir.",
                                   "Düzensizlikten rahatsızlık duymayabilir.",
                                   "Birden fazla işi eş zamanlı yapma tamamlayabilme eğilimleri düşüktür."]

    envanter_degerlendirme_19_0 = ["Hedeflerine odaklanabilir.", "Zorluklar karşısında dayanıklıdır.",
                                   "Duruma uygun davranışlar sergileyebilen bir yapıya sahip olabilir."]
    envanter_degerlendirme_19_1 = ["Sosyal ortamlarda kolayca incinebilen hassas bir yapıya sahip olabilir.",
                                   "Çevreleri tarafından onay alma ihtiyacı duyabilir.",
                                   "Kötümser bakış açısına sahip olabilir.",
                                   "Olaylar karşısında kendini suçlamaya yatkın olabilir."]

    envanter_degerlendirme_20_0 = ["Daha çok somut faaliyetlerden ve rutin işlerden hoşlanabilir.",
                                   "Geleneksel düşünce yapısına sahip olabilir.",
                                   "Yeni durumları tehdit olarak görebilir.",
                                   "Duygularıyla karar vermeyen nesnel bakış açısına sahip olabilir."]
    envanter_degerlendirme_20_1 = ["Gözlem yeteneği gelişmiş olabilir.",
                                   "İlgi alanı geniş olabilir ve olaylara değişik bakış açısı sergilyebilir. Kendinin ve başkasının duygularını anlamakta ustadır. Düşünceli, nazik, hassas olabilir ve yoğun duygu durumlarını yaşayabilir."]

    envanter_degerlendirme_21_0 = ["Sosyal ortamlarda neşeli, canlı, hareketli ve sıcak olmaya yatkın olabilir.",
                                   "Sosyal ortamlarda aktif olmaktan hoşlnabilir.",
                                   "Duygularıyla hareket etmeye yatkındır.",
                                   "Olumsuz tepkilerini içlerinde bastıramayabilir.",
                                   "Bu nedenle fevri davranışlar sergileyebilir."]
    envanter_degerlendirme_21_1 = ["Sosyal ortamlarda daha sessiz kalmayı ve gözlemci olmayı tercih edebilir.",
                                   "Duygularını kolayca dışa vurmayı tercih etmeyebilir.",
                                   "Daha ciddi ve kontrollü davranışlar sergileyebilir.",
                                   "Olumsuz duygularını dışa vurmaktansa içlerinde yaşamayı tercih edebilir."]

    envanter_degerlendirme_22_0 = ["Öz disiplin seviyesi düşük olabilir.",
                                   "Motivasyonu dışsal faktörlere bağlı olabilir.",
                                   "Organizasyon yeteneği düşük olabilir.",
                                   "Plan ile süreci birlikte yürütme sıkıntısı yaşayabilir."]
    envanter_degerlendirme_22_1 = ["Olaylar karşısında daha nesnel davranabilir.",
                                   "Amaçlarına ulaşmaya odaklanabilir, bu nedenle özdisiplinleri yüksek olabilir.",
                                   "Sorumluluklarını yerine getirmede içsel motivasyonları yüksektir olabilir."]

    envanter_degerlendirme_23_0 = ["Motivasyon seviyeleri yüksek, azimli kararlı bir şekilde çalışmaya yatkındır.",
                                   "Sorumluluklarını yerine getirmekte kararlıdır."]
    envanter_degerlendirme_23_1 = ["Kişisel hedef belirlemede zorluklar yaşabilir.",
                                   "Zamanını verimli şekilde kullanamayabilir.", "Motivasyonları düşmeye eğilimlidir."]

    envanter_degerlendirme_24_0 = ["Duygu durumunu genellikle kontrol edebilir.",
                                   "Davranışları nisbeten tahmin edilebilir.",
                                   "Tutarlı davranışlar sergileyebilen bir yapıya sahip olabilir."]
    envanter_degerlendirme_24_1 = ["Duygu durumu kolayca değişebilir.", "Davranışlarında tutarsızlıklar olabilir.",
                                   "Gergin ortamlarda motivasyonları düşebilir."]

    envanter_degerlendirme_25_0 = ["Entelektüel faaliyetlere pek ilgi göstermeyebilir.",
                                   "Fikirlerle fazla uğraşmayan daha çok somut düşünen bir yapıya sahip olabilir.",
                                   "Rutin faliyetlerde bulunmaya yatkındır."]
    envanter_degerlendirme_25_1 = ["Olayları irdeleyen, soyutlamalar yapan, kendini geliştirmeye açık olabilir.",
                                   "Olaylar karşısında daha geniş bir bakış açısına sahip olabilir.",
                                   "Akademik faaliyetlerden hoşlanan bir yapıya sahip olabilir."]

    envanter_degerlendirme_26_0 = ["Çevreden gelen bildirimlere hassas olabilir.",
                                   "Olumsuzluklar karşısında duygu durumu kolayca değişebilir.",
                                   "Olayları kişiselleştirebilir.", "Sorumluluk almayı sevmeyebilir.",
                                   "Hedefine ulaşmakta güçlükler yaşayabilir.",
                                   "Problemler karşısında sorunlar yaşayabilir."]
    envanter_degerlendirme_26_1 = ["Özgüveni yüksek olabilir.", "Kendinden emin olabilirler.",
                                   "Duygu durumu kolayca değişmez.",
                                   "Problemler karşısında etkili başa çıkma stratejilerine sahip olabilir."]

    envanter_degerlendirme_27_0 = ["İnsanlarla etkileşim halinde olmaktan hoşlanabilir.",
                                   "Bireysel çalışma yerine grupla çalışmalarda daha başarılı olabilir.",
                                   "Kararlarını alırken başkalarından etkilenebilir.",
                                   "Dışsal motivasyonu daha etkilidir."]
    envanter_degerlendirme_27_1 = ["Sosyal ortamlarda insanlarla iletişime geçmekten hoşlanmayabilir.",
                                   "Sosyal ortamlara girmektense yalnız kalmayı tercih edebilir.",
                                   "İçsel motivasyonu ve otokontrolü yüksek olabilir.",
                                   "Kararlarını verirken başkalarından fazla etkilenmeyebilir.",
                                   "Grup çalışması yerine bireysel çalışmalarda daha başarılı olabilir."]

    envanter_degerlendirme_28_0 = ["Hedefine ulaşamak için kendilerini organize etmede güçlükler yaşayabilir.",
                                   "Zaman yönetimi konusunda sorun yaşayabilir.",
                                   "Çabuk sıkılan bir yapıya sahip olabilir.",
                                   "Yaptığı işlerle ilgili fikirleri sıklıkla değişebilir."]
    envanter_degerlendirme_28_1 = ["Yüksek başarı motivasyonuna sahip olabilir.",
                                   "Hedeflerine ulaşmak için azimli ve aktif bir şekilde çalışabilir.",
                                   "Güvenilir ve sorumluluk sahibi olmaya yatkındır."]

    envanter_degerlendirme_29_0 = ["Duygu durumu kolayca değişmeyebilir.", "Tutarlı davranışlar sergileyebilir.",
                                   "Çok hassas bir yapıya sahip olmayabilir.", "Olaylardan kolay etkilenemeyebilir.",
                                   "Davranışları nisbeten önceden kestirilebilir.", "Güçlü bir yapıya sahip olabilir."]
    envanter_degerlendirme_29_1 = ["Duygu durumu kolayca değişebilir.",
                                   "Anlık duruma göre duruma göre duygusal tutarsızlıklar yaşabilir.",
                                   "Hassas bir yapıya sahip olabilir.", "Kolayca incinebilir.",
                                   "Olaylardan kolay etkilenebilir.",
                                   "Kararsız ve kendilerini suçlamaya yatkın olabilir."]

    envanter_degerlendirme_30_0 = ["Başkalarından ziyade kendine eğilimi yüksek olabilir.",
                                   "Olaylar karşısında nesnel davranışlar sergilemeye yatkın olabilir.",
                                   "Duygusal olmayan, gerçekçi ve kapalı bir düşünce yapısına sahip olabilir."]
    envanter_degerlendirme_30_1 = ["Duygulu, hassas, ince bir yapıya sahip olabilir.",
                                   "Duygularını yoğun şekilde yaşayabilir.", "Empati seviyeleri yüksektir.",
                                   "Estetik ve sanata yatkın olabilir."]

    envanter_degerlendirme_31_0 = ["Sosyal ortamlara kolayca uyum sağlayabilir.",
                                   "Neşeli,canlı,cesur ve rekabet hisleri yüksek olabilir.",
                                   "İçinde bulundukları grubu yönlendirmeyi ve grupta ön plana çıkmayı tercih edebilir."]
    envanter_degerlendirme_31_1 = ["Sosyal ortamlarda ön plana çıkmayı tercih etmeyebilir.",
                                   "Kalbalık ortamlardansa yalnız kalmayı tercih edebilir.",
                                   "Kalabalık ortamlarda kaygı yaşayabilir.",
                                   "İlk kez girdikleri ortamlarda sessiz, sakin ve resmi davranabilir.",
                                   "İlk kez girdikleri ortamlarda sessiz, sakin ve resmi davranabilir.",
                                   "Olumsuz duygularını içlerinde yaşamayı tercih edebilir."]

    envanter_degerlendirme_32_0 = ["Eleştirilmekten hoşlanmayabilir.",
                                   "Başkalarına karşı hoşgörü gösterme konusunda güçlük yaşayabilir.",
                                   "Kin tutmaya meyilli olabilir.", "Kibirli davranışlar sergileyebilir."]
    envanter_degerlendirme_32_1 = ["Yeni girdikleri ortama kolayca uyum sağlayabilir.",
                                   "Alçakgönüllü, pozitif ve girişken yapıya sahip olabilir.",
                                   "Sorunlar karşısında pozitif tavır sergileyebilir."]

    envanter_degerlendirme_33_0 = ["Hedeflerine ulaşma konusunda sıkıntılar yaşayabilir.",
                                   "Sorumluluklarını ertelemeye yatkın olabilir.",
                                   "Zamanlarını verimli kullanamayabilir.",
                                   "Organizasyon konusunda sıkıntı yaşayabilir."]
    envanter_degerlendirme_33_1 = ["Amaçlarına ulaşmak için gerekli içsel motivasyonları yüksek olabilir.",
                                   "Hedefler koyabilen ve gerçekleştirmek için çaba harcayan bir yapıya sahip olabilir."]

    envanter_degerlendirme_34_0 = ["Duygu durumları kolayca değişebilir.", "Anlık ruh haline göre davranabilir.",
                                   "Bu nedenle problemlerin çözümünde sakin kalamayabilir.",
                                   "Tutarsız, değişen duygusal yapı gösterebilir."]
    envanter_degerlendirme_34_1 = ["Olaylar karşısında mantıklı çözümler bulabilir.", "Olaylara objektif bakabilir.",
                                   "Probleri büyütmekten çok çözmeye odaklı olabilir.",
                                   "Davranışlar nisbeten önceden kestirilebilir.", "Tutarlı kişiliğe sahip olabilir."]

    envanter_degerlendirme_35_0 = ["Olaylara farklı yaklaşım tarzı sergileyebilir.",
                                   "Rutin işlerden sıkılmaya yatkın olabilir.",
                                   "Farklı deneyimler yaşamaya eğilimli olabilir.", "Enerji seviyesi yüksek olabilir."]
    envanter_degerlendirme_35_1 = ["Yetkin olduğu işleri yapmayı tercih edebilir.",
                                   "Çatışma ortamından kaçınmayı tercih edebilir.",
                                   "Olaylara farklı yaklaşımlar sergileyemeyebilir.", "Belirsizliği sevmeyebilir.",
                                   "Her yeniliği bir risk olarak görebilir."]

    envanter_degerlendirme_36_0 = ["Kalabalık ortamlardan çok sakin ortamları tercih edebilir.",
                                   "Yeni girdikleri sosyal ortamlarda sessiz, sakin daha çok gözlemci olabilir.",
                                   "Sır tutma konusunda güvenilir olabilir.",
                                   "Olumsuz tepkilerini içlerinde yaşamayı tercih edebilir."]
    envanter_degerlendirme_36_1 = ["Sosyal ortamlarda konuşkan, neşeli, hevesli ve canlı olabilir.",
                                   "Enerjileri yüksek olup doğal davranışlar sergileyebilir.",
                                   "Emir alma konusunda sıkıntılar yaşabilir.",
                                   "Tepki ve  düşüncelerini olumsuzda olsada paylaşmaya yatkın olabilir."]

    envanter_degerlendirme_37_0 = ["Sosyal ortamlarda ılımlı, anlayışlı tavırlar sergiler.",
                                   "Güven verici kişilik özelliklerine sahip olabilir.",
                                   "Peşin hükümlü bir yapıya sahip değildir bu nedenle eleştriye tahhamül seviyesi yüksek olabilir.",
                                   "Eleştrilmeye açık olabilir.", "Önyargısız ve hoşgörülü bir yapıya sahip olabilir."]
    envanter_degerlendirme_37_1 = ["Çabuk sinirlenebilen, önyargılı bir yapıya sahip olabilir.",
                                   "Peşin hükümlü bir yapıya sahip olduklarından olumsuz tepkilerini kolayca ortaya koyabilir.",
                                   "Eleştrilmekten hoşlanmayabilir.", "Zaman zaman insanlara karşı kırıcı olabilir."]

    envanter_degerlendirme_38_0 = ["İçsel motivasyonu düşük olabilir.",
                                   "İş organizasyonu konusunda sıkıntılar yaşabilir.",
                                   "Zorlandıklarını düşündükleri işlerde ertelemeye eğilimli olabilir.",
                                   "Otokontrol ve sorumluluk bilinçleri düşük olabilir."]
    envanter_degerlendirme_38_1 = ["İçsel motivasyonu yüksek olabilir.",
                                   "İşine kolayca odaklanabilen ve amaca yönelik hareket edebilen bir yapıya sahip olabilir.",
                                   "Detaylara dikkat edebilir.",
                                   "Hatasız iş yapmaya özen gösteren bir yapıya sahip olabilir."]

    envanter_degerlendirme_39_0 = ["Sosyal ortamlarda pozitif olmaya yatkındır.",
                                   "Durumu uygun davranışlar sergileyebilir.",
                                   "Sorgulayan ve güvenilir bir yapıya sahip olabilir."]
    envanter_degerlendirme_39_1 = ["Duygu durumu sıkça değişiklik gösterebilir.",
                                   "Tutarsız davranışlar sergileyebilir.",
                                   "Gerginliklerinden dolayı iş motivasyonları düşebilir.",
                                   "Sosyalleşme ve iletişim problemleri yaşayabilir."]

    envanter_degerlendirme_40_0 = ["Fikirlerle uğraşmaktansa somut faaliyetlerden daha fazla zevk alabilir.",
                                   "Entelektüel faaliyetlere pek ilgi göstermeyebilir."]
    envanter_degerlendirme_40_1 = ["Analitik düşünce yapısına sahip olabilir.", "Analiz yapma yetisi yüksektir.",
                                   "Soyut düşünebilir.", "Kendini geliştirmeye eğilimli bir yapıya sahip olabilir.",
                                   "Durumlara farklı açılardan yaklaşabilir."]

    envanter_degerlendirme_41_0 = ["Duygulu, ince ve sanatçı ruhlu bir yapıya sahip olabilir.",
                                   "Kendinin ve başkalarının duygularını önemseme eğilimi yüksek olabilir."]
    envanter_degerlendirme_41_1 = ["Nesnel düşünme eğilimi yüksek olabilir.",
                                   "Kendi avantajlarını daha fazla önemseyebilir.",
                                   "Düşünceleri duygusallıktan uzak olabilir.",
                                   "Karar alırken duygularından çok mantığını ön planda tutabilir.",
                                   "Hayal güçleri sınırlı olabilir."]

    envanter_degerlendirme_42_0 = ["Grup çalışmasından çok bireysel çalışmayı tercih edebilir.",
                                   "Rekabetçi bir kişilik yapısına sahip olabilir.",
                                   "Kendisini diğer insanlara göre daha fazla önemseyebilir.",
                                   "Başkalarına karşı anlayış seviyeleri düşük olabilir."]
    envanter_degerlendirme_42_1 = ["Başkalarını düşünen, empati düzeyi yüksek bir yapıya sahip olabilir.",
                                   "Merhametli, yardımsever, sıcak ve anlayışlı bir kişiliğe sahip olabilir.",
                                   "İşbirliğine yatkın ve insanlarla iyi geçinen bir yapıya sahip olabilir.",
                                   "Grup çalışmasına daha yatkın olabilir."]

    envanter_degerlendirme_43_0 = ["Planlı çalışma eğilimi yüksek olabilir.", "Konsantrasyonu yüksek olabilir.",
                                   "Otokontrolü yüksek olabilir.", "Başladığı işi bitirme eğilimi yüksek olabilir."]
    envanter_degerlendirme_43_1 = ["Otokontrol seviyesi düşük olabilir.", "Bir işi yaparken çok çabuk sıkılabilir.",
                                   "Detayları gözden kaçırmaya eğilimli olabilir."]

    envanter_degerlendirme_44_0 = ["Hayal gücü sınırlı, gerçi bir yapıya sahip olabilir.",
                                   "Mantıklı hareket etme eğilimi yüksek olabilir.", "Fikirleri sabit olabilir.",
                                   "Düşünceleri kolay kolay değişmez önyargılı olabilir."]
    envanter_degerlendirme_44_1 = ["Sanatsal ve estetik bakış açısı gelişmiş olabilir.",
                                   "Duyarlı bir yapıya sahip olabilir ve duygularını yoğun yaşayabilir.",
                                   "Olaylara farklı yönden bakabilir.",
                                   "İnce ruhlu, naif ve barışçıl bir yapıya sahip olabilir.",
                                   "Uzamsal zekası gelişmiş olabilir."]

    soru = int(ogrenci[19].values[0])
    if soru <= 3:
        envanter_analiz += envanter_degerlendirme_1_0[rd.randint(0, len(envanter_degerlendirme_1_0) - 1)]
    elif soru >= 6:
        envanter_analiz += envanter_degerlendirme_1_1[rd.randint(0, len(envanter_degerlendirme_1_1) - 1)]
    else:
        arada_kalanlar.append("1")

    soru = int(ogrenci[20].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_2_0[rd.randint(0, len(envanter_degerlendirme_2_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_2_1[rd.randint(0, len(envanter_degerlendirme_2_1) - 1)]
    else:
        arada_kalanlar.append("2")

    soru = int(ogrenci[21].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_3_0[rd.randint(0, len(envanter_degerlendirme_3_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_3_1[rd.randint(0, len(envanter_degerlendirme_3_1) - 1)]
    else:
        arada_kalanlar.append("3")

    soru = int(ogrenci[22].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_4_0[rd.randint(0, len(envanter_degerlendirme_4_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_4_1[rd.randint(0, len(envanter_degerlendirme_4_1) - 1)]
    else:
        arada_kalanlar.append("4")

    soru = int(ogrenci[23].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_5_0[rd.randint(0, len(envanter_degerlendirme_5_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_5_1[rd.randint(0, len(envanter_degerlendirme_5_1) - 1)]
    else:
        arada_kalanlar.append("5")

    soru = int(ogrenci[24].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_6_0[rd.randint(0, len(envanter_degerlendirme_6_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_6_1[rd.randint(0, len(envanter_degerlendirme_6_1) - 1)]
    else:
        arada_kalanlar.append("6")

    soru = int(ogrenci[25].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_7_0[rd.randint(0, len(envanter_degerlendirme_7_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_7_1[rd.randint(0, len(envanter_degerlendirme_7_1) - 1)]
    else:
        arada_kalanlar.append("7")

    soru = int(ogrenci[26].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_8_0[rd.randint(0, len(envanter_degerlendirme_8_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_8_1[rd.randint(0, len(envanter_degerlendirme_8_1) - 1)]
    else:
        arada_kalanlar.append("8")

    soru = int(ogrenci[27].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_9_0[rd.randint(0, len(envanter_degerlendirme_9_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_9_1[rd.randint(0, len(envanter_degerlendirme_9_1) - 1)]
    else:
        arada_kalanlar.append("9")

    soru = int(ogrenci[28].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_10_0[rd.randint(0, len(envanter_degerlendirme_10_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_10_1[rd.randint(0, len(envanter_degerlendirme_10_1) - 1)]
    else:
        arada_kalanlar.append("10")

    soru = int(ogrenci[29].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_11_0[rd.randint(0, len(envanter_degerlendirme_11_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_11_1[rd.randint(0, len(envanter_degerlendirme_11_1) - 1)]
    else:
        arada_kalanlar.append("11")

    soru = int(ogrenci[30].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_12_0[rd.randint(0, len(envanter_degerlendirme_12_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_12_1[rd.randint(0, len(envanter_degerlendirme_12_1) - 1)]
    else:
        arada_kalanlar.append("12")

    soru = int(ogrenci[31].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_13_0[rd.randint(0, len(envanter_degerlendirme_13_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_13_1[rd.randint(0, len(envanter_degerlendirme_13_1) - 1)]
    else:
        arada_kalanlar.append("13")

    soru = int(ogrenci[32].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_14_0[rd.randint(0, len(envanter_degerlendirme_14_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_14_1[rd.randint(0, len(envanter_degerlendirme_14_1) - 1)]
    else:
        arada_kalanlar.append("14")

    soru = int(ogrenci[33].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_15_0[rd.randint(0, len(envanter_degerlendirme_15_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_15_1[rd.randint(0, len(envanter_degerlendirme_15_1) - 1)]
    else:
        arada_kalanlar.append("15")

    soru = int(ogrenci[34].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_16_0[rd.randint(0, len(envanter_degerlendirme_16_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_16_1[rd.randint(0, len(envanter_degerlendirme_16_1) - 1)]
    else:
        arada_kalanlar.append("16")

    soru = int(ogrenci[35].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_17_0[rd.randint(0, len(envanter_degerlendirme_17_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_17_1[rd.randint(0, len(envanter_degerlendirme_17_1) - 1)]
    else:
        arada_kalanlar.append("17")

    soru = int(ogrenci[36].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_18_0[rd.randint(0, len(envanter_degerlendirme_18_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_18_1[rd.randint(0, len(envanter_degerlendirme_18_1) - 1)]
    else:
        arada_kalanlar.append("18")

    soru = int(ogrenci[37].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_19_0[rd.randint(0, len(envanter_degerlendirme_19_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_19_1[rd.randint(0, len(envanter_degerlendirme_19_1) - 1)]
    else:
        arada_kalanlar.append("19")

    soru = int(ogrenci[38].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_20_0[rd.randint(0, len(envanter_degerlendirme_20_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_20_1[rd.randint(0, len(envanter_degerlendirme_20_1) - 1)]
    else:
        arada_kalanlar.append("20")

    soru = int(ogrenci[39].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_21_0[rd.randint(0, len(envanter_degerlendirme_21_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_21_1[rd.randint(0, len(envanter_degerlendirme_21_1) - 1)]
    else:
        arada_kalanlar.append("21")

    soru = int(ogrenci[40].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_22_0[rd.randint(0, len(envanter_degerlendirme_22_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_22_1[rd.randint(0, len(envanter_degerlendirme_22_1) - 1)]
    else:
        arada_kalanlar.append("22")

    soru = int(ogrenci[41].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_23_0[rd.randint(0, len(envanter_degerlendirme_23_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_23_1[rd.randint(0, len(envanter_degerlendirme_23_1) - 1)]
    else:
        arada_kalanlar.append("23")

    soru = int(ogrenci[42].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_24_0[rd.randint(0, len(envanter_degerlendirme_24_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_24_1[rd.randint(0, len(envanter_degerlendirme_24_1) - 1)]
    else:
        arada_kalanlar.append("24")

    soru = int(ogrenci[43].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_25_0[rd.randint(0, len(envanter_degerlendirme_25_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_25_1[rd.randint(0, len(envanter_degerlendirme_25_1) - 1)]
    else:
        arada_kalanlar.append("25")

    soru = int(ogrenci[44].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_26_0[rd.randint(0, len(envanter_degerlendirme_26_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_26_1[rd.randint(0, len(envanter_degerlendirme_26_1) - 1)]
    else:
        arada_kalanlar.append("26")

    soru = int(ogrenci[45].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_27_0[rd.randint(0, len(envanter_degerlendirme_27_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_27_1[rd.randint(0, len(envanter_degerlendirme_27_1) - 1)]
    else:
        arada_kalanlar.append("27")

    soru = int(ogrenci[46].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_28_0[rd.randint(0, len(envanter_degerlendirme_28_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_28_1[rd.randint(0, len(envanter_degerlendirme_28_1) - 1)]
    else:
        arada_kalanlar.append("28")

    soru = int(ogrenci[47].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_29_0[rd.randint(0, len(envanter_degerlendirme_29_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_29_1[rd.randint(0, len(envanter_degerlendirme_29_1) - 1)]
    else:
        arada_kalanlar.append("29")

    soru = int(ogrenci[48].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_30_0[rd.randint(0, len(envanter_degerlendirme_30_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_30_1[rd.randint(0, len(envanter_degerlendirme_30_1) - 1)]
    else:
        arada_kalanlar.append("30")

    soru = int(ogrenci[49].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_31_0[rd.randint(0, len(envanter_degerlendirme_31_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_31_1[rd.randint(0, len(envanter_degerlendirme_31_1) - 1)]
    else:
        arada_kalanlar.append("31")

    soru = int(ogrenci[50].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_32_0[rd.randint(0, len(envanter_degerlendirme_32_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_32_1[rd.randint(0, len(envanter_degerlendirme_32_1) - 1)]
    else:
        arada_kalanlar.append("32")

    soru = int(ogrenci[51].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_33_0[rd.randint(0, len(envanter_degerlendirme_33_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_33_1[rd.randint(0, len(envanter_degerlendirme_33_1) - 1)]
    else:
        arada_kalanlar.append("33")

    soru = int(ogrenci[52].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_34_0[rd.randint(0, len(envanter_degerlendirme_34_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_34_1[rd.randint(0, len(envanter_degerlendirme_34_1) - 1)]
    else:
        arada_kalanlar.append("34")

    soru = int(ogrenci[53].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_35_0[rd.randint(0, len(envanter_degerlendirme_35_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_35_1[rd.randint(0, len(envanter_degerlendirme_35_1) - 1)]
    else:
        arada_kalanlar.append("35")

    soru = int(ogrenci[54].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_36_0[rd.randint(0, len(envanter_degerlendirme_36_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_36_1[rd.randint(0, len(envanter_degerlendirme_36_1) - 1)]
    else:
        arada_kalanlar.append("36")

    soru = int(ogrenci[55].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_37_0[rd.randint(0, len(envanter_degerlendirme_37_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_37_1[rd.randint(0, len(envanter_degerlendirme_37_1) - 1)]
    else:
        arada_kalanlar.append("37")

    soru = int(ogrenci[56].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_38_0[rd.randint(0, len(envanter_degerlendirme_38_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_38_1[rd.randint(0, len(envanter_degerlendirme_38_1) - 1)]
    else:
        arada_kalanlar.append("38")

    soru = int(ogrenci[57].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_39_0[rd.randint(0, len(envanter_degerlendirme_39_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_39_1[rd.randint(0, len(envanter_degerlendirme_39_1) - 1)]
    else:
        arada_kalanlar.append("39")

    soru = int(ogrenci[58].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_40_0[rd.randint(0, len(envanter_degerlendirme_40_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_40_1[rd.randint(0, len(envanter_degerlendirme_40_1) - 1)]
    else:
        arada_kalanlar.append("40")

    soru = int(ogrenci[59].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_41_0[rd.randint(0, len(envanter_degerlendirme_41_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_41_1[rd.randint(0, len(envanter_degerlendirme_41_1) - 1)]
    else:
        arada_kalanlar.append("41")

    soru = int(ogrenci[60].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_42_0[rd.randint(0, len(envanter_degerlendirme_42_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_42_1[rd.randint(0, len(envanter_degerlendirme_42_1) - 1)]
    else:
        arada_kalanlar.append("42")

    soru = int(ogrenci[61].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_43_0[rd.randint(0, len(envanter_degerlendirme_43_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_43_1[rd.randint(0, len(envanter_degerlendirme_43_1) - 1)]
    else:
        arada_kalanlar.append("43")

    soru = int(ogrenci[62].values[0])
    if soru <= 3:
        envanter_analiz += " " + envanter_degerlendirme_44_0[rd.randint(0, len(envanter_degerlendirme_44_0) - 1)]
    elif soru >= 6:
        envanter_analiz += " " + envanter_degerlendirme_44_1[rd.randint(0, len(envanter_degerlendirme_44_1) - 1)]
    else:
        arada_kalanlar.append("44")

    return envanter_analiz, arada_kalanlar


def ogrenci_analiz_olustur(ogrenci):
    cevaplar = (
        ("1 --> " + str(ogrenci[19].values[0]), "12 --> " + str(ogrenci[30].values[0]),
         "23 --> " + str(ogrenci[41].values[0]), "34 --> " + str(ogrenci[52].values[0])),
        ("2 --> " + str(ogrenci[20].values[0]), "13 --> " + str(ogrenci[31].values[0]),
         "24 --> " + str(ogrenci[42].values[0]), "35 --> " + str(ogrenci[53].values[0])),
        ("3 --> " + str(ogrenci[21].values[0]), "14 --> " + str(ogrenci[32].values[0]),
         "25 --> " + str(ogrenci[43].values[0]), "36 --> " + str(ogrenci[54].values[0])),
        ("4 --> " + str(ogrenci[22].values[0]), "15 --> " + str(ogrenci[33].values[0]),
         "26 --> " + str(ogrenci[44].values[0]), "37 --> " + str(ogrenci[55].values[0])),
        ("5 --> " + str(ogrenci[23].values[0]), "16 --> " + str(ogrenci[34].values[0]),
         "27 --> " + str(ogrenci[45].values[0]), "38 --> " + str(ogrenci[56].values[0])),
        ("6 --> " + str(ogrenci[24].values[0]), "17 --> " + str(ogrenci[35].values[0]),
         "28 --> " + str(ogrenci[46].values[0]), "39 --> " + str(ogrenci[57].values[0])),
        ("7 --> " + str(ogrenci[25].values[0]), "18 --> " + str(ogrenci[36].values[0]),
         "29 --> " + str(ogrenci[47].values[0]), "40 --> " + str(ogrenci[58].values[0])),
        ("8 --> " + str(ogrenci[26].values[0]), "19 --> " + str(ogrenci[37].values[0]),
         "30 --> " + str(ogrenci[48].values[0]), "41 --> " + str(ogrenci[59].values[0])),
        ("9 --> " + str(ogrenci[27].values[0]), "20 --> " + str(ogrenci[38].values[0]),
         "31 --> " + str(ogrenci[49].values[0]), "42 --> " + str(ogrenci[60].values[0])),
        ("10 --> " + str(ogrenci[28].values[0]), "21 --> " + str(ogrenci[39].values[0]),
         "32 --> " + str(ogrenci[50].values[0]), "43 --> " + str(ogrenci[61].values[0])),
        ("11 --> " + str(ogrenci[29].values[0]), "22 --> " + str(ogrenci[40].values[0]),
         "33 --> " + str(ogrenci[51].values[0]), "44 --> " + str(ogrenci[62].values[0]))
    )

    ogr_no = str(ogrenci[2].values[0]).title()
    boyut_a = (int(ogrenci[19].values[0]) + (11 - (int(ogrenci[24].values[0]))) + int(ogrenci[29].values[0]) + int(
        ogrenci[34].values[0]) + (11 - int(ogrenci[39].values[0])) + int(ogrenci[44].values[0]) + (
                           11 - int(ogrenci[49].values[0])) + int(ogrenci[54].values[0])) / 8
    boyut_b = ((11 - int(ogrenci[20].values[0])) + int(ogrenci[25].values[0]) + (11 - int(ogrenci[30].values[0])) + int(
        ogrenci[35].values[0]) + int(ogrenci[40].values[0]) + (11 - int(ogrenci[45].values[0])) + int(
        ogrenci[50].values[0]) + (11 - int(ogrenci[55].values[0])) + int(ogrenci[60].values[0])) / 9
    boyut_c = (int(ogrenci[21].values[0]) + int(ogrenci[31].values[0]) + int(ogrenci[46].values[0]) + int(
        ogrenci[51].values[0]) + int(ogrenci[56].values[0]) + (11 - int(ogrenci[26].values[0])) + (
                           11 - int(ogrenci[36].values[0])) + (11 - int(ogrenci[41].values[0])) + (
                           11 - int(ogrenci[61].values[0]))) / 9
    boyut_d = (int(ogrenci[22].values[0]) + int(ogrenci[32].values[0]) + int(ogrenci[37].values[0]) + int(
        ogrenci[47].values[0]) + int(ogrenci[57].values[0]) + (11 - int(ogrenci[27].values[0])) + (
                           11 - int(ogrenci[41].values[0])) + (11 - int(ogrenci[52].values[0]))) / 8
    boyut_e = (int(ogrenci[23].values[0]) + int(ogrenci[28].values[0]) + int(ogrenci[33].values[0]) + int(
        ogrenci[38].values[0]) + int(ogrenci[43].values[0]) + int(ogrenci[48].values[0]) + int(
        ogrenci[58].values[0]) + int(ogrenci[62].values[0]) + (11 - int(ogrenci[53].values[0])) + (
                           11 - int(ogrenci[59].values[0]))) / 10
    a = '{0:.2f}'.format(boyut_a)
    b = '{0:.2f}'.format(boyut_b)
    c = '{0:.2f}'.format(boyut_c)
    d = '{0:.2f}'.format(boyut_d)
    e = '{0:.2f}'.format(boyut_e)

    document = Document()

    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "\tSakarya Üniversitesi Kariyer ve Yetenek Yönetimi Koordinatörlüğü"
    paragraph.style = document.styles["Header"]

    document.add_heading(str(ogrenci[2].values[0]).title(), 0)

    p = document.add_paragraph()
    p.add_run('Envanter Doldurulma Tarihi: ').bold = True
    p.add_run(str(ogrenci[0].values[0])).italic = True

    p = document.add_paragraph()
    analiz, arada_kalanlar = tanımla_analiz_et(ogrenci)
    p.add_run('Envanter Analizi').bold = True
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    analiz_uzunlugu = int(len(analiz) * 0.4)
    analiz = ' '.join(analiz.split()[:analiz_uzunlugu]) + "..."
    paragraph.add_run(analiz).italic = True

    p = document.add_paragraph()
    p.add_run('Detaylı analiz için lütfen kariyer@sakarya.edu.tr adresiyle iletişime geçin.').bold = True


    dosya_adi = str(ogrenci[2].values[0]).title() + ".docx"
    document.save(dosya_adi)


def danisman_analiz_olustur(ogrenci):
    cevaplar = (
        ("1 --> " + str(ogrenci[19].values[0]), "12 --> " + str(ogrenci[30].values[0]),
         "23 --> " + str(ogrenci[41].values[0]), "34 --> " + str(ogrenci[52].values[0])),
        ("2 --> " + str(ogrenci[20].values[0]), "13 --> " + str(ogrenci[31].values[0]),
         "24 --> " + str(ogrenci[42].values[0]), "35 --> " + str(ogrenci[53].values[0])),
        ("3 --> " + str(ogrenci[21].values[0]), "14 --> " + str(ogrenci[32].values[0]),
         "25 --> " + str(ogrenci[43].values[0]), "36 --> " + str(ogrenci[54].values[0])),
        ("4 --> " + str(ogrenci[22].values[0]), "15 --> " + str(ogrenci[33].values[0]),
         "26 --> " + str(ogrenci[44].values[0]), "37 --> " + str(ogrenci[55].values[0])),
        ("5 --> " + str(ogrenci[23].values[0]), "16 --> " + str(ogrenci[34].values[0]),
         "27 --> " + str(ogrenci[45].values[0]), "38 --> " + str(ogrenci[56].values[0])),
        ("6 --> " + str(ogrenci[24].values[0]), "17 --> " + str(ogrenci[35].values[0]),
         "28 --> " + str(ogrenci[46].values[0]), "39 --> " + str(ogrenci[57].values[0])),
        ("7 --> " + str(ogrenci[25].values[0]), "18 --> " + str(ogrenci[36].values[0]),
         "29 --> " + str(ogrenci[47].values[0]), "40 --> " + str(ogrenci[58].values[0])),
        ("8 --> " + str(ogrenci[26].values[0]), "19 --> " + str(ogrenci[37].values[0]),
         "30 --> " + str(ogrenci[48].values[0]), "41 --> " + str(ogrenci[59].values[0])),
        ("9 --> " + str(ogrenci[27].values[0]), "20 --> " + str(ogrenci[38].values[0]),
         "31 --> " + str(ogrenci[49].values[0]), "42 --> " + str(ogrenci[60].values[0])),
        ("10 --> " + str(ogrenci[28].values[0]), "21 --> " + str(ogrenci[39].values[0]),
         "32 --> " + str(ogrenci[50].values[0]), "43 --> " + str(ogrenci[61].values[0])),
        ("11 --> " + str(ogrenci[29].values[0]), "22 --> " + str(ogrenci[40].values[0]),
         "33 --> " + str(ogrenci[51].values[0]), "44 --> " + str(ogrenci[62].values[0]))
    )

    ogr_no = str(ogrenci[2].values[0]).title()
    boyut_a = (int(ogrenci[19].values[0]) + (11 - (int(ogrenci[24].values[0]))) + int(ogrenci[29].values[0]) + int(
        ogrenci[34].values[0]) + (11 - int(ogrenci[39].values[0])) + int(ogrenci[44].values[0]) + (
                           11 - int(ogrenci[49].values[0])) + int(ogrenci[54].values[0])) / 8
    boyut_b = ((11 - int(ogrenci[20].values[0])) + int(ogrenci[25].values[0]) + (11 - int(ogrenci[30].values[0])) + int(
        ogrenci[35].values[0]) + int(ogrenci[40].values[0]) + (11 - int(ogrenci[45].values[0])) + int(
        ogrenci[50].values[0]) + (11 - int(ogrenci[55].values[0])) + int(ogrenci[60].values[0])) / 9
    boyut_c = (int(ogrenci[21].values[0]) + int(ogrenci[31].values[0]) + int(ogrenci[46].values[0]) + int(
        ogrenci[51].values[0]) + int(ogrenci[56].values[0]) + (11 - int(ogrenci[26].values[0])) + (
                           11 - int(ogrenci[36].values[0])) + (11 - int(ogrenci[41].values[0])) + (
                           11 - int(ogrenci[61].values[0]))) / 9
    boyut_d = (int(ogrenci[22].values[0]) + int(ogrenci[32].values[0]) + int(ogrenci[37].values[0]) + int(
        ogrenci[47].values[0]) + int(ogrenci[57].values[0]) + (11 - int(ogrenci[27].values[0])) + (
                           11 - int(ogrenci[41].values[0])) + (11 - int(ogrenci[52].values[0]))) / 8
    boyut_e = (int(ogrenci[23].values[0]) + int(ogrenci[28].values[0]) + int(ogrenci[33].values[0]) + int(
        ogrenci[38].values[0]) + int(ogrenci[43].values[0]) + int(ogrenci[48].values[0]) + int(
        ogrenci[58].values[0]) + int(ogrenci[62].values[0]) + (11 - int(ogrenci[53].values[0])) + (
                           11 - int(ogrenci[59].values[0]))) / 10
    a = '{0:.2f}'.format(boyut_a)
    b = '{0:.2f}'.format(boyut_b)
    c = '{0:.2f}'.format(boyut_c)
    d = '{0:.2f}'.format(boyut_d)
    e = '{0:.2f}'.format(boyut_e)

    document = Document()

    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "\tSakarya Üniversitesi Kariyer ve Yetenek Yönetimi Koordinatörlüğü"
    paragraph.style = document.styles["Header"]

    document.add_heading(str(ogrenci[2].values[0]).title(), 0)

    p = document.add_paragraph()
    p.add_run('Envanter Doldurulma Tarihi: ').bold = True
    p.add_run(str(ogrenci[0].values[0])).italic = True

    p = document.add_paragraph()
    p.add_run('Mail adresi: ').bold = True
    p.add_run(str(ogrenci[1].values[0])).italic = True

    p = document.add_paragraph()
    p.add_run('Cinsiyeti: ').bold = True
    p.add_run(str(ogrenci[5].values[0])).italic = True

    p = document.add_paragraph()
    p.add_run('Yaşı: ').bold = True
    p.add_run(str(ogrenci[6].values[0])).italic = True

    p = document.add_paragraph()
    p.add_run('Bölüme giriş sırası: ').bold = True
    p.add_run(str(ogrenci[9].values[0])).italic = True

    p = document.add_paragraph()
    p.add_run('Bölümü sınava kaçıncı girişinde kazandı: ').bold = True
    p.add_run(str(ogrenci[10].values[0])).italic = True

    p = document.add_paragraph()
    p.add_run('Bölüm kaçıncı tercihiydi: ').bold = True
    p.add_run(str(ogrenci[11].values[0])).italic = True

    p = document.add_paragraph()
    p.add_run('Bölümü neden tercih etti: ').bold = True
    p.add_run(str(ogrenci[12].values[0])).italic = True

    p = document.add_paragraph()
    p.add_run('Kariyer Kavramı sizin için ne ifade etmektedir: ').bold = True
    p.add_run("1-" + str(ogrenci[13].values[0])).italic = True
    p.add_run(" 2-" + str(ogrenci[14].values[0])).italic = True
    p.add_run(" 3-" + str(ogrenci[15].values[0])).italic = True

    p = document.add_paragraph()
    p.add_run('Çalışmak istediği işte aradığı temel özellik: ').bold = True
    p.add_run(str(ogrenci[16].values[0])).italic = True

    p = document.add_paragraph()
    p.add_run('Faktörler: ').bold = True
    p.add_run(str(ogrenci[17].values[0])).italic = True

    p = document.add_paragraph()
    p.add_run('Kariyer yapılması planlanan sektör: ').bold = True
    p.add_run(str(ogrenci[18].values[0])).italic = True

    p = document.add_paragraph()
    analiz, arada_kalanlar = tanımla_analiz_et(ogrenci)
    p.add_run('Envanter Analizi').bold = True
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run(analiz).italic = True

    if len(arada_kalanlar) > 0:
        arada_kalan_mesaj = ""
        for i in arada_kalanlar[:-1]:
            arada_kalan_mesaj += i + " "
        arada_kalan_mesaj += "ve " + arada_kalanlar[len(arada_kalanlar) - 1] + " numaralı sorulara 4 ve 5 verilmiş."
        p = document.add_paragraph()
        p.add_run('Arada kalınan sorular: ').bold = True
        p.add_run(arada_kalan_mesaj).italic = True

    p = document.add_paragraph()
    p.add_run('Verilen Cevaplar').bold = True
    table = document.add_table(rows=1, cols=4)
    table.allow_autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Soru No --> Cevap'
    hdr_cells[1].text = 'Soru No --> Cevap'
    hdr_cells[2].text = 'Soru No --> Cevap'
    hdr_cells[3].text = 'Soru No --> Cevap'
    for bir, iki, uc, dort in cevaplar:
        row = table.add_row().cells
        row[0].text = bir
        row[1].text = iki
        row[2].text = uc
        row[3].text = dort
    table.style = 'Light Shading Accent 4'

    p = document.add_paragraph()
    p = document.add_paragraph()
    p.add_run('Boyut Analizi').bold = True
    table = document.add_table(rows=1, cols=6)
    table.allow_autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Öğrenci Adı'
    hdr_cells[1].text = 'Dışa Dönüklük'
    hdr_cells[2].text = 'Yumuşak Başlılık'
    hdr_cells[3].text = 'Özdenetim'
    hdr_cells[4].text = 'Duygusal Tutarlılık'
    hdr_cells[5].text = 'Gelişime Açıklık'
    row = table.add_row().cells
    row[0].text = str(ogr_no)
    row[1].text = a + "/10"
    row[2].text = b + "/10"
    row[3].text = c + "/10"
    row[4].text = d + "/10"
    row[5].text = e + "/10"
    table.style = 'Light Shading Accent 4'

    dosya_adi = (str(ogrenci[2].values[0]).title()).rstrip() + " Analiz.docx"
    document.save(dosya_adi)


def mime_init(from_addr, recipients_addr, subject, body):
    """
    :param str from_addr:           The email address you want to send mail from
    :param list recipients_addr:    The list of email addresses of recipients
    :param str subject:             Mail subject
    :param str body:                Mail body
    :return:                        MIMEMultipart object
    """
    msg = MIMEMultipart()
    msg['From'] = from_addr
    msg['To'] = ','.join(recipients_addr)
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))
    return msg


def send_email(user, password, from_addr, recipients_addr, subject, body, files_path=None, server='smtp.gmail.com'):
    """
    :param str user:                Sender's email userID
    :param str password:            sender's email password
    :param str from_addr:           The email address you want to send mail from
    :param list recipients_addr:    List of (or space separated string) email addresses of recipients
    :param str subject:             Mail subject
    :param str body:                Mail body
    :param list files_path:         List of paths of files you want to attach
    :param str server:              SMTP server (port is choosen 587)
    :return:                        None
    """

    #   assert isinstance(recipents_addr, list)
    FROM = from_addr
    TO = recipients_addr if isinstance(recipients_addr, list) else recipients_addr.split(' ')
    PASS = password
    SERVER = server
    SUBJECT = subject
    BODY = body
    msg = mime_init(FROM, TO, SUBJECT, BODY)

    for file_path in files_path or []:
        with open(file_path, "rb") as fp:
            part = MIMEBase('application', "octet-stream")
            part.set_payload((fp).read())
            # Encoding payload is necessary if encoded (compressed) file has to be attached.
            encoders.encode_base64(part)
            part.add_header('content-disposition', 'attachment', filename='%s' % file_path)
            msg.attach(part)

    if SERVER == 'localhost':  # send mail from local server
        # Start local SMTP server
        server = smtplib.SMTP(SERVER)
        text = msg.as_string()
        server.send_message(msg)
    else:
        # Start SMTP server at port 587
        server = smtplib.SMTP(SERVER, 587)
        server.starttls()
        # Enter login credentials for the email you want to sent mail from
        server.login(user, PASS)
        text = msg.as_string()
        # Send mail
        server.sendmail(FROM, TO, text)

    server.quit()


def mail_gonder(ogr_adi, ogr_maili):
    file_path = []
    user = 'kariyer@sakarya.edu.tr'  # Email userID
    password = st.secrets['kariyer_sifre']  # Email password
    from_addr = 'kariyer@sakarya.edu.tr'
    recipients_addr = ogr_maili
    subject = '5FE Kişilik Envanteri Analizi'
    body = "Sayın {}, Kariyer ve Yetenek Yönetimi Koordinatörlüğü'nde tamamlamış olduğunuz kişilik envanterinin analizi ektedir.".format(
        ogr_adi)
    file_path.append(ogr_adi + ".docx")
    # print(recipients_addr,file_path)
    send_email(user, password, from_addr, recipients_addr, subject, body, file_path)


def mail_gonder_yetkili(ogr_adi):
    file_path = []
    user = 'kariyer@sakarya.edu.tr'  # Email userID
    password = st.secrets['kariyer_sifre']  # Email password
    from_addr = 'kariyer@sakarya.edu.tr'
    recipients_addr = "kariyer@sakarya.edu.tr"
    subject = '5FE Kişilik Envanteri Analizi - {}'.format(ogr_adi)
    body = "{} isimli öğrenciye ait kişilik envanterinin analizi ektedir.".format(
        ogr_adi)
    file_path.append(ogr_adi.rstrip() + " Analiz.docx")
    send_email(user, password, from_addr, recipients_addr, subject, body, file_path)


def head():
    st.markdown("""
        <h1 style='text-align: center'>
        5FE Kişilik Envanteri Analizi
        </h1>
    """, unsafe_allow_html=True
                )

    st.caption("""
        <p style='text-align: center;'>
        Sakarya Üniversitesi Kariyer ve Yetenek Yönetimi Koordinatörlüğü
        </p>
    """, unsafe_allow_html=True
               )

def versiyon():
    st.caption("""
                <p style='text-align: center;'>
                ver 1.2.0<br/><font size="2">build 26062023.2222</font>
                </p>
            """, unsafe_allow_html=True
               )

